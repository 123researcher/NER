from itertools import groupby
import win32com.client as win32

# All the same as yours
word = win32.Dispatch("Word.Application")
word.Visible = 0
word.Documents.Open("G:/AUDIT REPORTS/fy21/Audit Report 21-05 - GTO BBPO Client Data Protection.docx")
doc = word.ActiveDocument

# Here we use itertools.groupby (without sorting anything) to
# find groups of words that share the same heading (note it picks
# up newlines). The tricky/confusing thing here is that you can't
# just group on the Style itself - you have to group on the str(). 
# There was some other interesting behavior, but I have zero 
# experience with COMObjects so I'll leave it there :)
# All of these comments for two lines of code :)
for heading, grp_wrds in groupby(doc.Words, key=lambda x: str(x.Style)):
    
  
 

    print(heading,"".join(str(word) for word in grp_wrds))
    
    
    
import docx2txt
import re
my_text = docx2txt.process("")


m = re.search('(?<=To)(\S+\s+){2}(\w+)',my_text)
#user = re.search('To\s\w+\s(\w+)',s)

print(m.group())


import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
print(getText("a.docx"))






import glob
import os

os.chdir(r'G:\AUDIT REPORTS\fy21\PDF')
myFiles = glob.glob('*.pdf')
print(myFiles)






import pandas as pd
writer = pd.ExcelWriter('output.xlsx', engine='xlsxwriter')
writer.save()
import pandas as pd

# dataframe columns
df = pd.DataFrame({'Audit Report Title': ['A', 'B', 'C', 'D'],'Email Address':['A', 'B', 'C', 'D'], 'Report Relationship':['A', 'B', 'C', 'D],
'

                   'Total Clicks': [10, 0, 30, 50],'Total Opens': [1, 0, 0, 40]})
                                                                                                                           
                
                                                                                                                                                                                                                            

# Create a Pandas Excel writer using XlsxWriter as the engine.
writer = pd.ExcelWriter('output.xlsx', engine='xlsxwriter')

# Convert the dataframe to an XlsxWriter Excel object.
df.to_excel(writer, sheet_name='Sheet1', index=False)

# Close the Pandas Excel writer and output the Excel file.
writer.save()

import pandas as pd
reader = pd.read_excel(r'output.xlsx')
print(reader)





import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)
print(getText("a.docx"))








